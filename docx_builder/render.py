"""Render a DocumentContent into a DOCX using the Jinja-instrumented template.

Usage:
    from docx_builder.render import render
    render('output/template_jinja.docx', doc_content, 'output/output.docx',
           image_base='output')

Requires:  pip install docxtpl python-docx lxml
"""
from __future__ import annotations
import collections
import copy
import os
import re
import subprocess
import tempfile
from pathlib import Path

_RE_HAS_INLINE_MATH = re.compile(r'\$[^$\n]+\$')
_RE_SECREF_CELL = re.compile(r':ref\{([^}]+)\}:')

from docx import Document  # type: ignore
from docx.shared import Cm  # type: ignore
from docxtpl import DocxTemplate, InlineImage, RichText  # type: ignore
from lxml import etree  # type: ignore

from .schema import ContentNode, Run, finalize_content

# XML namespace URIs used in image post-processing
_NS_W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
_NS_A  = 'http://schemas.openxmlformats.org/drawingml/2006/main'


# ---------------------------------------------------------------------------
# Structured-VAR image strings
# ---------------------------------------------------------------------------

# A structured VAR value (TYPE:yaml, TYPE:table cell, plain VAR) that is
# exactly a markdown image — ![alt](path) with an optional [width] hint —
# renders as a native inline image instead of text.
_RE_IMG_STRING = re.compile(r'^!\[([^\]]*)\]\(([^)\s]+)\)\s*(?:\[([^\]]+)\])?$')


def _width_from_hint(hint: str, text_width: int) -> int:
    """Width in EMU from a '50%' / '7cm' hint; full text width otherwise."""
    if hint.endswith('%'):
        try:
            return int(text_width * float(hint[:-1]) / 100)
        except ValueError:
            return text_width
    if hint.endswith('cm'):
        try:
            return Cm(float(hint[:-2]))
        except ValueError:
            return text_width
    return text_width


def _convert_context_images(value, tpl: DocxTemplate, img_base: Path, text_width: int):
    """Recursively replace markdown-image strings in a context value with
    InlineImage objects (and YAML nulls with empty strings)."""
    if value is None:
        return ''
    if isinstance(value, str):
        m = _RE_IMG_STRING.match(value.strip())
        if not m:
            return value
        path = m.group(2)
        width = _width_from_hint(m.group(3) or '', text_width)
        full = Path(path) if Path(path).is_absolute() else img_base / path
        if full.exists():
            return InlineImage(tpl, str(full), width=width)
        return RichText(f'[image not found: {path}]', italic=True)
    if isinstance(value, list):
        return [_convert_context_images(v, tpl, img_base, text_width) for v in value]
    if isinstance(value, dict):
        return {k: _convert_context_images(v, tpl, img_base, text_width)
                for k, v in value.items()}
    return value


# ---------------------------------------------------------------------------
# Run list → RichText
# ---------------------------------------------------------------------------

def _to_rt(runs: list[Run]) -> RichText:
    rt = RichText()
    for run in runs:
        kwargs = dict(bold=run.bold, italic=run.italic)
        if run.code:
            kwargs['font'] = 'Courier New'
        if run.color:
            kwargs['color'] = run.color
        rt.add(run.text, **kwargs)
    return rt


def _to_text(runs: list[Run]) -> str:
    return ''.join(r.text for r in runs)


def _runs_with_section_label(node: ContentNode) -> list[Run]:
    """Return node runs, appending heading section-label sentinel when present."""
    out = list(node.runs)
    if node.sec_id:
        out.append(Run(text=f'@@SECLABEL:{node.sec_id}@@'))
    return out


# ---------------------------------------------------------------------------
# Proxy classes — wrap ContentNode instances for Jinja template access
# ---------------------------------------------------------------------------

class HeadingProxy:
    """Proxy for h1–h4, paragraph, and bullet nodes.

    NOTE: docxtpl RichText objects do not render inside {%p if %} blocks.
    We therefore render a unique placeholder string and restore formatted runs
    in _post_process_inline_runs().
    """

    def __init__(
        self,
        node: ContentNode,
        math_counter: int = 0,
        text_placeholder: str | None = None,
        comment_id: int = 0,
    ) -> None:
        self._node = node
        self.type = node.type
        self._math_counter = math_counter
        self._text_placeholder = text_placeholder
        # Non-zero only for a math paragraph carrying a native Word comment;
        # _post_process_math() brackets the OMML paragraph with this id.
        self.comment_id = comment_id

    @property
    def text(self) -> str:
        if self._text_placeholder:
            return self._text_placeholder
        t = _to_text(self._node.runs)
        if self._node.sec_id:
            t += f'@@SECLABEL:{self._node.sec_id}@@'
        return t

    @property
    def has_math(self) -> bool:
        return self._node.has_math

    @property
    def math_placeholder(self) -> str:
        return f'MATH_PARA_{self._math_counter}'


class BulletProxy(HeadingProxy):
    """Same as HeadingProxy; bullet style is applied by the template paragraph."""
    @property
    def level(self) -> int:
        return self._node.level


class MathDisplayProxy:
    """Proxy for display math nodes ($$...$$).

    Renders a sentinel string in the template; _post_process_math() replaces it
    with the actual OMML paragraph generated by pandoc.
    """

    def __init__(self, node: ContentNode, counter: int) -> None:
        self._node = node
        self.type = 'math_display'
        self.math_placeholder = f'MATH_DISP_{counter}'

    # has_math is not needed for display math nodes, but keep interface consistent
    has_math: bool = False


class FigureProxy:
    """Proxy for figure nodes.

    Attributes used in the template:
      item.type         — 'figure'
      item.fig_number   — sequential figure number
      item.image        — InlineImage (or placeholder RichText)
      item.title        — plain string  "Figure N: Short Title"  (bold label)
      item.caption      — plain string  detailed description paragraph
      item.caption_text — plain string for LOF entries (short title only)
    """

    def __init__(self, node: ContentNode, tpl: DocxTemplate, image_base: Path, text_width: int) -> None:
        self._node = node
        self._tpl = tpl
        self._image_base = image_base
        self.type = 'figure'
        self._text_width = text_width
        self.fig_number = node.fig_number
        # LOF uses the short title, not the long description
        self.caption_text = _to_text(node.fig_title) or _to_text(node.fig_caption)
        # Cache caption text and detect inline math for post-processing
        self._caption_text = _to_text(node.fig_caption)
        self.has_caption_math = bool(_RE_HAS_INLINE_MATH.search(self._caption_text)) if self._caption_text else False
        self.caption_math_placeholder = f'MATH_CAP_{node.fig_number}' if self.has_caption_math else ''

    def _resolve_width(self) -> int:
        """Return image width in EMU from fig_width_hint, or full text width."""
        hint = self._node.fig_width_hint
        if not hint:
            return self._text_width
        if hint.endswith('%'):
            try:
                pct = float(hint[:-1])
                return int(self._text_width * pct / 100)
            except ValueError:
                return self._text_width
        if hint.endswith('cm'):
            try:
                cm = float(hint[:-2])
                return Cm(cm)
            except ValueError:
                return self._text_width
        return self._text_width

    @property
    def image(self) -> InlineImage | RichText:
        path = self._node.fig_path
        if not path:
            return RichText(f'[Figure {self.fig_number}: no image]', italic=True)

        width = self._resolve_width()

        if path.startswith('_mermaid:'):
            fig_id = path[len('_mermaid:'):]
            # Look for a pre-rendered PNG next to the draft
            for candidate in [
                self._image_base / f'{fig_id}.png',
                self._image_base / 'figures' / f'{fig_id}.png',
                self._image_base / 'snaps' / f'{fig_id}.png',
            ]:
                if candidate.exists():
                    return InlineImage(self._tpl, str(candidate), width=width)
            return RichText(
                f'[Figure {self.fig_number}: Mermaid diagram — run mmdc to render {fig_id}.png]',
                italic=True,
            )

        # Resolve relative path
        full = Path(path) if Path(path).is_absolute() else self._image_base / path
        if full.exists():
            return InlineImage(self._tpl, str(full), width=width)
        return RichText(f'[Figure {self.fig_number}: {path} not found]', italic=True)

    @property
    def title(self) -> str:
        """Short label: 'Figure N: Short Title' — shown above the image."""
        short = _to_text(self._node.fig_title)
        return f'Figure {self.fig_number}: {short}' if short else ''

    @property
    def caption(self) -> RichText | str:
        """Detailed description — italic text only, shown below the figure-title label.

        When the caption contains inline math, returns a sentinel string
        (MATH_CAP_N) that _post_process_math() will replace with OMML.
        """
        if self.has_caption_math:
            return self.caption_math_placeholder
        if not self._caption_text:
            return RichText('')
        rt = RichText()
        rt.add(self._caption_text, italic=True)
        return rt

    @property
    def fig_start(self) -> str:
        return f'FIGURE_START_{self.fig_number}'

    @property
    def fig_end(self) -> str:
        return f'FIGURE_END_{self.fig_number}'


class TableProxy:
    """Proxy for table nodes.

    Attributes used in the template:
      item.type         — 'table'
      item.tbl_number   — sequential table number
      item.caption      — RichText  "Table N: caption text"
      item.caption_text — plain string for LOT entries
      item.subdoc       — SubDocxTemplate containing the rendered table
    """

    def __init__(self, node: ContentNode, tpl: DocxTemplate) -> None:
        self._node = node
        self._tpl = tpl
        self.type = 'table'
        self.tbl_number = node.tbl_number
        self.caption_text = _to_text(node.tbl_caption)

    @property
    def caption(self) -> str:
        return f'Table {self.tbl_number}: {_to_text(self._node.tbl_caption)}'

    @property
    def subdoc(self) -> str:
        """Return a unique placeholder string.

        Subdocs inside {%p if %} blocks do not work with docxtpl.
        We use a sentinel string; _post_process_tables() replaces it with
        the actual table after docxtpl rendering.
        """
        return f'TABLE_PLACEHOLDER_{self.tbl_number}'


# ---------------------------------------------------------------------------
# Proxy factory
# ---------------------------------------------------------------------------

def _build_content_proxies(
    content: list[ContentNode],
    tpl: DocxTemplate,
    image_base: Path,
    text_width: int,
) -> tuple[list, dict[str, list[Run]], list[dict]]:
    """Build proxy list, assigning sequential math counters."""
    proxies = []
    inline_map: dict[str, list[Run]] = {}
    comments: list[dict] = []
    math_disp_n = 0
    math_para_n = 0
    inline_n = 0
    comment_n = 0

    def _register_comment(node: ContentNode, preplaced: bool = False) -> int:
        """Append node's comment to the comments list; return its id (0 if none).

        `preplaced=True` marks a comment whose range markers are inserted
        directly by a later pass (math paragraphs, see _post_process_math),
        rather than via @@COMMENTSTART/END@@ sentinel runs.
        """
        nonlocal comment_n
        if not node.comment_author:
            return 0
        comment_n += 1
        comments.append({
            'id': comment_n,
            'author': node.comment_author,
            'body': node.comment_body,
            'preplaced': preplaced,
        })
        return comment_n

    def _wrap_comment(node: ContentNode, runs: list[Run]) -> list[Run]:
        cid = _register_comment(node)
        if not cid:
            return runs
        return (
            [Run(text=f'@@COMMENTSTART:{cid}@@')]
            + runs
            + [Run(text=f'@@COMMENTEND:{cid}@@')]
        )

    for node in content:
        if node.comment_author and node.type in {'figure', 'table', 'math_display'}:
            raise ValueError(
                f"Comment by {node.comment_author!r} attached to an "
                f"unsupported {node.type!r} node — this should have been "
                f"rejected during parsing."
            )
        if node.type == 'figure':
            proxies.append(FigureProxy(node, tpl, image_base, text_width))
        elif node.type == 'table':
            proxies.append(TableProxy(node, tpl))
        elif node.type == 'math_display':
            math_disp_n += 1
            proxies.append(MathDisplayProxy(node, math_disp_n))
        elif node.type == 'bullet' and node.has_math:
            math_para_n += 1
            proxies.append(BulletProxy(
                node, math_para_n,
                comment_id=_register_comment(node, preplaced=True)))
        elif node.type in {'p', 'numbered'} and node.has_math:
            math_para_n += 1
            proxies.append(HeadingProxy(
                node, math_counter=math_para_n,
                comment_id=_register_comment(node, preplaced=True)))
        elif node.type in {'h1', 'h2', 'h3', 'h4', 'p', 'numbered'}:
            inline_n += 1
            key = f'@@INLINEFMT:{inline_n}@@'
            inline_map[key] = _wrap_comment(node, _runs_with_section_label(node))
            proxies.append(HeadingProxy(node, text_placeholder=key))
        elif node.type == 'bullet':
            inline_n += 1
            key = f'@@INLINEFMT:{inline_n}@@'
            inline_map[key] = _wrap_comment(node, _runs_with_section_label(node))
            proxies.append(BulletProxy(node, text_placeholder=key))
        else:
            # Includes '_footnote_def' nodes: no template branch matches their
            # type, so they render no body content — kept 1:1 so the per-VAR
            # proxy split in render() stays aligned.
            proxies.append(HeadingProxy(node))
    return proxies, inline_map, comments


# ---------------------------------------------------------------------------
# Footnotes — registry construction and validation
# ---------------------------------------------------------------------------

def _collect_footnotes(content_nodes: list[ContentNode]) -> list[dict]:
    """Build the footnote list ({key, runs} in first-reference order) from
    '_footnote_def' nodes and @@FOOTREF:key@@ reference runs.

    Raises on duplicate definitions, undefined references, and references in
    unsupported places (figure titles/captions, table captions/cells).
    Warns (stdout) about definitions that are never referenced.
    """
    defs: dict[str, list[Run]] = {}
    for node in content_nodes:
        if node.type != '_footnote_def':
            continue
        if node.footnote_key in defs:
            raise ValueError(f"Footnote [^{node.footnote_key}] is defined more than once.")
        defs[node.footnote_key] = node.runs

    key_order: list[str] = []
    for node in content_nodes:
        if node.type == '_footnote_def':
            continue
        if node.type == 'figure':
            if any(r.footnote_key for r in node.fig_title + node.fig_caption):
                raise ValueError(
                    f"Figure {node.fig_id!r} has a footnote reference in its "
                    f"title/caption — footnotes are only supported in body "
                    f"text (headings, paragraphs, list items)."
                )
            continue
        if node.type == 'table':
            if any(r.footnote_key for r in node.tbl_caption) or any(
                '@@FOOTREF:' in cell
                for row in [node.tbl_headers, *node.tbl_rows] for cell in row
            ):
                raise ValueError(
                    f"Table {node.tbl_id!r} has a footnote reference in its "
                    f"caption or cells — footnotes are only supported in body "
                    f"text (headings, paragraphs, list items)."
                )
            continue
        for run in node.runs:
            if run.footnote_key and run.footnote_key not in key_order:
                key_order.append(run.footnote_key)

    missing = [k for k in key_order if k not in defs]
    if missing:
        raise ValueError(
            "Footnote reference(s) without a definition: "
            + ", ".join(f"[^{k}]" for k in missing)
        )
    unused = [k for k in defs if k not in key_order]
    if unused:
        print("Warning: footnote definition(s) never referenced: "
              + ", ".join(f"[^{k}]" for k in unused))

    return [{'key': k, 'runs': defs[k]} for k in key_order]


def _post_process_inline_runs(docx_path: str, inline_map: dict[str, list[Run]]) -> None:
    """Replace inline-format placeholders with explicit formatted Word runs."""
    if not inline_map:
        return

    from docx import Document as DocxDoc
    from docx.oxml import OxmlElement

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    def _set_bold(r_pr):
        for tag in ('b', 'bCs'):
            if r_pr.find(f'{{{W}}}{tag}') is None:
                r_pr.append(OxmlElement(f'w:{tag}'))

    def _set_italic(r_pr):
        for tag in ('i', 'iCs'):
            if r_pr.find(f'{{{W}}}{tag}') is None:
                r_pr.append(OxmlElement(f'w:{tag}'))

    def _set_color(r_pr, color_val):
        color_el = OxmlElement('w:color')
        # Normalize: ensure hex colors have no '#' prefix for OOXML
        val = color_val.lstrip('#')
        color_el.set(f'{{{W}}}val', val)
        if r_pr.find(f'{{{W}}}color') is None:
            r_pr.append(color_el)

    def _append_run(p_el, run: Run):
        if not run.text:
            return
        r_el = OxmlElement('w:r')
        if run.bold or run.italic or run.code or run.color:
            r_pr = OxmlElement('w:rPr')
            if run.bold:
                _set_bold(r_pr)
            if run.italic:
                _set_italic(r_pr)
            if run.code:
                fonts_el = OxmlElement('w:rFonts')
                fonts_el.set(f'{{{W}}}ascii', 'Courier New')
                fonts_el.set(f'{{{W}}}hAnsi', 'Courier New')
                r_pr.append(fonts_el)
            if run.color:
                _set_color(r_pr, run.color)
            r_el.append(r_pr)
        t_el = OxmlElement('w:t')
        t_el.set(XML_SPACE, 'preserve')
        t_el.text = run.text
        r_el.append(t_el)
        p_el.append(r_el)

    doc = DocxDoc(docx_path)
    replaced = 0

    all_paras = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        key = para.text.strip()
        runs = inline_map.get(key)
        if runs is None:
            continue

        p_el = para._p
        for r in list(p_el.findall(f'{{{W}}}r')):
            p_el.remove(r)

        for run in runs:
            _append_run(p_el, run)

        replaced += 1

    if replaced:
        doc.save(docx_path)
        print(f'Post-processed {replaced} inline-formatted paragraph(s) in {docx_path}')


# ---------------------------------------------------------------------------
# Public render function
# ---------------------------------------------------------------------------

def _post_process_tables(docx_path: str, tables_list: list[ContentNode]) -> None:
    """Replace TABLE_PLACEHOLDER_N paragraphs with actual python-docx tables."""
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = DocxDocument(docx_path)

    # Build a lookup: tbl_number → ContentNode
    tbl_map = {n.tbl_number: n for n in tables_list}

    replaced = 0
    for para in list(doc.paragraphs):
        txt = para.text.strip()
        if not txt.startswith('TABLE_PLACEHOLDER_'):
            continue
        try:
            tbl_num = int(txt.split('TABLE_PLACEHOLDER_')[1])
        except (IndexError, ValueError):
            continue
        node = tbl_map.get(tbl_num)
        if node is None:
            continue

        headers = node.tbl_headers
        rows = node.tbl_rows
        if not headers:
            continue
        n_cols = len(headers)
        n_rows = 1 + len(rows)

        # Build the table (appended to end initially)
        tbl = doc.add_table(rows=n_rows, cols=n_cols, style='Table Grid')

        # Header row — bold
        for ci, hdr in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = hdr
            for run in cell.paragraphs[0].runs:
                run.bold = True

        # Data rows
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data[:n_cols]):
                val = _RE_SECREF_CELL.sub(lambda m: f'@@SECREF:{m.group(1)}@@', val)
                tbl.rows[ri + 1].cells[ci].text = val

        # Move table to immediately after the placeholder paragraph
        para._p.addnext(tbl._tbl)

        # Remove placeholder paragraph
        para._p.getparent().remove(para._p)
        replaced += 1

    if replaced:
        doc.save(docx_path)
        print(f'Post-processed {replaced} table(s) into {docx_path}')


_NS_W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'
_RE_CHK_SENTINEL = re.compile(r'@@CHK:([01])@@')


def _post_process_checkboxes(docx_path: str) -> None:
    """Sync native w14:checkbox controls from rendered @@CHK:x@@ sentinels.

    A prep script opts a checkbox in by replacing its content-run text with a
    Jinja expression that renders '@@CHK:1@@' or '@@CHK:0@@'
    (template_tools.jinjify_checkboxes). This pass sets w14:checked and the
    state glyph declared by the control itself — byte-for-byte what Word
    writes when a user clicks the box, so the control stays live. Documents
    with no sentinels are left untouched.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)
    changed = 0
    for sdt in doc.element.body.iter(f'{{{_NS_W}}}sdt'):
        checkbox = sdt.find(f'.//{{{_NS_W14}}}checkbox')
        content = sdt.find(f'{{{_NS_W}}}sdtContent')
        if checkbox is None or content is None:
            continue
        t_els = content.findall(f'.//{{{_NS_W}}}t')
        m = _RE_CHK_SENTINEL.fullmatch(''.join(t.text or '' for t in t_els).strip())
        if not m:
            continue
        checked = m.group(1) == '1'

        state_tag = 'checkedState' if checked else 'uncheckedState'
        state = checkbox.find(f'{{{_NS_W14}}}{state_tag}')
        glyph = '☒' if checked else '☐'
        if state is not None and state.get(f'{{{_NS_W14}}}val'):
            glyph = chr(int(state.get(f'{{{_NS_W14}}}val'), 16))

        checked_el = checkbox.find(f'{{{_NS_W14}}}checked')
        if checked_el is not None:
            checked_el.set(f'{{{_NS_W14}}}val', '1' if checked else '0')

        t_els[0].text = glyph
        for t in t_els[1:]:
            t.text = ''
        changed += 1

    if changed:
        doc.save(docx_path)
        print(f'Post-processed {changed} checkbox(es) in {docx_path}')


def _runs_to_w_elements(parent, runs: list[Run], xml_space: str) -> None:
    """Serialize Runs as <w:r> children of an lxml element (footnote/comment
    bodies: plain formatted text only)."""
    for run in runs:
        r_el = etree.SubElement(parent, f'{{{_NS_W}}}r')
        if run.bold or run.italic or run.code or run.color:
            r_pr = etree.SubElement(r_el, f'{{{_NS_W}}}rPr')
            if run.bold:
                etree.SubElement(r_pr, f'{{{_NS_W}}}b')
            if run.italic:
                etree.SubElement(r_pr, f'{{{_NS_W}}}i')
            if run.code:
                fonts_el = etree.SubElement(r_pr, f'{{{_NS_W}}}rFonts')
                fonts_el.set(f'{{{_NS_W}}}ascii', 'Courier New')
                fonts_el.set(f'{{{_NS_W}}}hAnsi', 'Courier New')
            if run.color:
                color_el = etree.SubElement(r_pr, f'{{{_NS_W}}}color')
                color_el.set(f'{{{_NS_W}}}val', run.color.lstrip('#'))
        t_el = etree.SubElement(r_el, f'{{{_NS_W}}}t')
        t_el.set(xml_space, 'preserve')
        t_el.text = run.text


def _post_process_footnotes(docx_path: str, footnotes: list[dict]) -> None:
    """Convert @@FOOTREF:key@@ sentinel runs into native Word footnotes:
    a FootnoteReference-styled w:footnoteReference run in document.xml plus
    one w:footnote entry per key appended to word/footnotes.xml.

    Word assigns the *displayed* numbers by reference order automatically —
    w:id values only bind references to entries, so new entries are keyed
    from max(existing id)+1 and coexist with any footnotes the template
    already carries. If the template has no footnotes part at all, one is
    created (with separator/continuation-separator entries) and registered
    in [Content_Types].xml and document.xml.rels — same raw-zip approach as
    _post_process_comments, and like it this must run after every pass that
    calls doc.save() on a Document opened before the zip patch.
    """
    if not footnotes:
        return

    import shutil
    import zipfile

    from docx import Document as DocxDoc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'
    RE_FOOTREF_SENT = re.compile(r'@@FOOTREF:([^@]+)@@')

    # --- Read (or scaffold) the footnotes part to allocate ids ---
    CT = 'http://schemas.openxmlformats.org/package/2006/content-types'
    PR = 'http://schemas.openxmlformats.org/package/2006/relationships'
    REL_TYPE = ('http://schemas.openxmlformats.org/officeDocument/2006/'
                'relationships/footnotes')
    FOOT_CT = ('application/vnd.openxmlformats-officedocument.'
               'wordprocessingml.footnotes+xml')
    rels_path = 'word/_rels/document.xml.rels'

    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = set(zin.namelist())
        foot_xml = zin.read('word/footnotes.xml') if 'word/footnotes.xml' in names else None

    part_existed = foot_xml is not None
    if part_existed:
        foot_root = etree.fromstring(foot_xml)
        existing_ids = []
        for el in foot_root.findall(f'{{{_NS_W}}}footnote'):
            try:
                existing_ids.append(int(el.get(f'{{{_NS_W}}}id') or ''))
            except ValueError:
                pass
        base_id = max(existing_ids + [0]) + 1
    else:
        foot_root = etree.Element(f'{{{_NS_W}}}footnotes', nsmap={'w': _NS_W})
        for sep_id, sep_tag, sep_type in (
            ('-1', 'separator', 'separator'),
            ('0', 'continuationSeparator', 'continuationSeparator'),
        ):
            fn = etree.SubElement(foot_root, f'{{{_NS_W}}}footnote')
            fn.set(f'{{{_NS_W}}}type', sep_type)
            fn.set(f'{{{_NS_W}}}id', sep_id)
            p_el = etree.SubElement(fn, f'{{{_NS_W}}}p')
            r_el = etree.SubElement(p_el, f'{{{_NS_W}}}r')
            etree.SubElement(r_el, f'{{{_NS_W}}}{sep_tag}')
        base_id = 1

    runs_by_key = {f['key']: f['runs'] for f in footnotes}

    # --- Replace sentinel runs in document.xml with footnote references ---
    # Word's model is one footnote entry per reference: a single entry
    # referenced from two places is malformed and triggers "Word found a
    # problem... we repaired it", after which the renumber drops a number.
    # So every @@FOOTREF@@ occurrence gets its own id and its own entry, even
    # when several share a key (the citation text is simply duplicated) — the
    # same thing Word does when you insert the same footnote twice by hand.
    doc = DocxDoc(docx_path)
    body = doc.element.body
    placed = 0
    next_id = base_id
    placements = []  # (id, key) in document order — one per reference
    for p_el in body.iter(qn('w:p')):
        for r_el in list(p_el.findall(qn('w:r'))):
            text = ''.join(t.text or '' for t in r_el.findall(qn('w:t')))
            if '@@FOOTREF:' not in text:
                continue
            new_els = []

            def _plain_run(fragment: str):
                run_el = copy.deepcopy(r_el)
                for child in list(run_el):
                    if child.tag != qn('w:rPr'):
                        run_el.remove(child)
                t_el = OxmlElement('w:t')
                t_el.set(XML_SPACE, 'preserve')
                t_el.text = fragment
                run_el.append(t_el)
                return run_el

            pos = 0
            for m in RE_FOOTREF_SENT.finditer(text):
                if m.start() > pos:
                    new_els.append(_plain_run(text[pos:m.start()]))
                key = m.group(1)
                if key not in runs_by_key:
                    raise ValueError(
                        f"Footnote reference [^{key}] has no definition "
                        f"(should have been caught by _collect_footnotes)."
                    )
                fid = next_id
                next_id += 1
                placements.append((fid, key))
                ref_run = OxmlElement('w:r')
                r_pr = OxmlElement('w:rPr')
                r_style = OxmlElement('w:rStyle')
                r_style.set(qn('w:val'), 'FootnoteReference')
                r_pr.append(r_style)
                ref_run.append(r_pr)
                ref_el = OxmlElement('w:footnoteReference')
                ref_el.set(qn('w:id'), str(fid))
                ref_run.append(ref_el)
                new_els.append(ref_run)
                placed += 1
                pos = m.end()
            if pos < len(text):
                new_els.append(_plain_run(text[pos:]))

            for el in new_els:
                r_el.addprevious(el)
            p_el.remove(r_el)

    if not placed:
        print('Warning: footnote definitions present but no @@FOOTREF@@ '
              'sentinels found in the document body — footnotes not written.')
        return

    # Every footnote id actually referenced in the body — both the references
    # we just placed and any the template already carried.
    referenced_ids = {
        ref.get(qn('w:id'))
        for ref in body.iter(qn('w:footnoteReference'))
        if ref.get(qn('w:id')) is not None
    }

    doc.save(docx_path)

    # --- Append the footnote entries to the footnotes part ---
    for fid, key in placements:
        fn = etree.SubElement(foot_root, f'{{{_NS_W}}}footnote')
        fn.set(f'{{{_NS_W}}}id', str(fid))
        p_el = etree.SubElement(fn, f'{{{_NS_W}}}p')
        p_pr = etree.SubElement(p_el, f'{{{_NS_W}}}pPr')
        p_style = etree.SubElement(p_pr, f'{{{_NS_W}}}pStyle')
        p_style.set(f'{{{_NS_W}}}val', 'FootnoteText')
        ref_r = etree.SubElement(p_el, f'{{{_NS_W}}}r')
        ref_rpr = etree.SubElement(ref_r, f'{{{_NS_W}}}rPr')
        ref_rstyle = etree.SubElement(ref_rpr, f'{{{_NS_W}}}rStyle')
        ref_rstyle.set(f'{{{_NS_W}}}val', 'FootnoteReference')
        etree.SubElement(ref_r, f'{{{_NS_W}}}footnoteRef')
        space_r = etree.SubElement(p_el, f'{{{_NS_W}}}r')
        space_t = etree.SubElement(space_r, f'{{{_NS_W}}}t')
        space_t.set(XML_SPACE, 'preserve')
        space_t.text = ' '
        _runs_to_w_elements(p_el, runs_by_key[key], XML_SPACE)

    # Drop orphan footnote entries the template shipped as authoring examples:
    # a real note (no w:type) that nothing in the body references. Left in
    # place, renderers surface it — Word/LibreOffice number it and shift every
    # citation by one; OnlyOffice overlaps it into the body text. Structural
    # entries (separator/continuationSeparator/continuationNotice, all typed)
    # and our own just-appended notes are all kept.
    for fn in list(foot_root.findall(f'{{{_NS_W}}}footnote')):
        if (fn.get(f'{{{_NS_W}}}type') is None
                and fn.get(f'{{{_NS_W}}}id') not in referenced_ids):
            foot_root.remove(fn)

    new_foot_xml = etree.tostring(foot_root, xml_declaration=True,
                                  encoding='UTF-8', standalone=True)

    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = set(zin.namelist())
        ct_xml = zin.read('[Content_Types].xml')
        rels_xml = zin.read(rels_path) if rels_path in names else None
        others = {n: zin.read(n) for n in names
                  if n not in ('[Content_Types].xml', rels_path, 'word/footnotes.xml')}

    ct_root = etree.fromstring(ct_xml)
    if not any(el.get('PartName') == '/word/footnotes.xml'
               for el in ct_root.findall(f'{{{CT}}}Override')):
        override = etree.SubElement(ct_root, f'{{{CT}}}Override')
        override.set('PartName', '/word/footnotes.xml')
        override.set('ContentType', FOOT_CT)
    new_ct_xml = etree.tostring(ct_root, xml_declaration=True,
                                encoding='UTF-8', standalone=True)

    if rels_xml is not None:
        rels_root = etree.fromstring(rels_xml)
    else:
        rels_root = etree.Element(f'{{{PR}}}Relationships', nsmap={None: PR})
    if not any(el.get('Type') == REL_TYPE
               for el in rels_root.findall(f'{{{PR}}}Relationship')):
        existing_rids = []
        for el in rels_root.findall(f'{{{PR}}}Relationship'):
            m = re.match(r'rId(\d+)$', el.get('Id') or '')
            if m:
                existing_rids.append(int(m.group(1)))
        rel = etree.SubElement(rels_root, f'{{{PR}}}Relationship')
        rel.set('Id', f'rId{(max(existing_rids) + 1) if existing_rids else 1}')
        rel.set('Type', REL_TYPE)
        rel.set('Target', 'footnotes.xml')
    new_rels_xml = etree.tostring(rels_root, xml_declaration=True,
                                  encoding='UTF-8', standalone=True)

    tmp_path = docx_path + '.tmp'
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in others.items():
            zout.writestr(name, data)
        zout.writestr('[Content_Types].xml', new_ct_xml)
        zout.writestr(rels_path, new_rels_xml)
        zout.writestr('word/footnotes.xml', new_foot_xml)
    shutil.move(tmp_path, docx_path)

    print(f'Post-processed {placed} footnote reference(s) '
          f'({len(placements)} footnote entries from {len(runs_by_key)} '
          f'unique definition(s)) into {docx_path}')


def _comment_initials(author: str) -> str:
    parts = author.split()
    return ''.join(p[0].upper() for p in parts if p) or '?'


def _post_process_comments(docx_path: str, comments: list[dict]) -> None:
    """Convert @@COMMENTSTART:n@@ / @@COMMENTEND:n@@ sentinel runs (emitted
    by _build_content_proxies, materialized by _post_process_inline_runs)
    into a real Word (OOXML) comment: commentRangeStart/End and a
    commentReference run in word/document.xml, plus the word/comments.xml
    part and its content-type/relationship registrations.

    Must run last: the package-level parts added below (comments.xml, the
    content-type override, the relationship) are written via a raw zip
    patch that a later doc.save() from another post-processing pass could
    silently drop, since python-docx doesn't know those parts exist.
    """
    if not comments:
        return

    from docx import Document as DocxDoc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    doc = DocxDoc(docx_path)
    body = doc.element.body

    def _find_run(sentinel: str):
        for r in body.iter(qn('w:r')):
            t = r.find(qn('w:t'))
            if t is not None and (t.text or '') == sentinel:
                return r
        return None

    placed = 0
    for c in comments:
        cid = c['id']
        if c.get('preplaced'):
            # Range markers already inserted by _post_process_math(); only the
            # comments.xml part and package wiring below remain to be written.
            placed += 1
            continue
        start_r = _find_run(f'@@COMMENTSTART:{cid}@@')
        end_r = _find_run(f'@@COMMENTEND:{cid}@@')
        if start_r is None or end_r is None:
            continue  # defensive: should not happen if md_parser validated correctly

        start_p = start_r.getparent()
        rng_start = OxmlElement('w:commentRangeStart')
        rng_start.set(qn('w:id'), str(cid))
        start_p.insert(list(start_p).index(start_r), rng_start)
        start_p.remove(start_r)

        end_p = end_r.getparent()
        rng_end = OxmlElement('w:commentRangeEnd')
        rng_end.set(qn('w:id'), str(cid))
        end_idx = list(end_p).index(end_r)
        end_p.insert(end_idx, rng_end)
        end_p.remove(end_r)

        ref_run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'CommentReference')
        r_pr.append(r_style)
        ref_run.append(r_pr)
        ref_el = OxmlElement('w:commentReference')
        ref_el.set(qn('w:id'), str(cid))
        ref_run.append(ref_el)
        end_p.insert(end_idx + 1, ref_run)

        placed += 1

    if not placed:
        return

    doc.save(docx_path)

    # --- Package-level parts python-docx doesn't manage ---
    import shutil
    import zipfile

    CT = 'http://schemas.openxmlformats.org/package/2006/content-types'
    PR = 'http://schemas.openxmlformats.org/package/2006/relationships'
    REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    rels_path = 'word/_rels/document.xml.rels'

    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = set(zin.namelist())
        ct_xml = zin.read('[Content_Types].xml')
        rels_xml = zin.read(rels_path) if rels_path in names else None
        others = {n: zin.read(n) for n in names
                  if n not in ('[Content_Types].xml', rels_path)}

    # word/comments.xml
    comments_root = etree.Element(f'{{{_NS_W}}}comments', nsmap={'w': _NS_W})
    for c in comments:
        w_comment = etree.SubElement(comments_root, f'{{{_NS_W}}}comment')
        w_comment.set(f'{{{_NS_W}}}id', str(c['id']))
        w_comment.set(f'{{{_NS_W}}}author', c['author'])
        w_comment.set(f'{{{_NS_W}}}initials', _comment_initials(c['author']))
        w_comment.set(f'{{{_NS_W}}}date', '2026-01-01T00:00:00Z')
        for para_runs in c['body']:
            p_el = etree.SubElement(w_comment, f'{{{_NS_W}}}p')
            for run in para_runs:
                r_el = etree.SubElement(p_el, f'{{{_NS_W}}}r')
                if run.bold or run.italic or run.code or run.color:
                    r_pr = etree.SubElement(r_el, f'{{{_NS_W}}}rPr')
                    if run.bold:
                        etree.SubElement(r_pr, f'{{{_NS_W}}}b')
                    if run.italic:
                        etree.SubElement(r_pr, f'{{{_NS_W}}}i')
                    if run.code:
                        fonts_el = etree.SubElement(r_pr, f'{{{_NS_W}}}rFonts')
                        fonts_el.set(f'{{{_NS_W}}}ascii', 'Courier New')
                        fonts_el.set(f'{{{_NS_W}}}hAnsi', 'Courier New')
                    if run.color:
                        color_el = etree.SubElement(r_pr, f'{{{_NS_W}}}color')
                        color_el.set(f'{{{_NS_W}}}val', run.color.lstrip('#'))
                t_el = etree.SubElement(r_el, f'{{{_NS_W}}}t')
                t_el.set(XML_SPACE, 'preserve')
                t_el.text = run.text
    comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # [Content_Types].xml — register comments.xml if not already present
    ct_root = etree.fromstring(ct_xml)
    if not any(el.get('PartName') == '/word/comments.xml'
               for el in ct_root.findall(f'{{{CT}}}Override')):
        override = etree.SubElement(ct_root, f'{{{CT}}}Override')
        override.set('PartName', '/word/comments.xml')
        override.set('ContentType',
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
    new_ct_xml = etree.tostring(ct_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # word/_rels/document.xml.rels — register the comments relationship
    if rels_xml is not None:
        rels_root = etree.fromstring(rels_xml)
    else:
        rels_root = etree.Element(f'{{{PR}}}Relationships', nsmap={None: PR})
    if not any(el.get('Type') == REL_TYPE for el in rels_root.findall(f'{{{PR}}}Relationship')):
        existing_ids = []
        for el in rels_root.findall(f'{{{PR}}}Relationship'):
            m = re.match(r'rId(\d+)$', el.get('Id') or '')
            if m:
                existing_ids.append(int(m.group(1)))
        rel = etree.SubElement(rels_root, f'{{{PR}}}Relationship')
        rel.set('Id', f'rId{(max(existing_ids) + 1) if existing_ids else 1}')
        rel.set('Type', REL_TYPE)
        rel.set('Target', 'comments.xml')
    new_rels_xml = etree.tostring(rels_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    tmp_path = docx_path + '.tmp'
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in others.items():
            zout.writestr(name, data)
        zout.writestr('[Content_Types].xml', new_ct_xml)
        zout.writestr(rels_path, new_rels_xml)
        zout.writestr('word/comments.xml', comments_xml)
    shutil.move(tmp_path, docx_path)

    print(f'Post-processed {placed} native Word comment(s) into {docx_path}')


_OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _pandoc_math_para(src: str, display: bool):
    """Convert LaTeX math (display or inline paragraph markdown) to an OMML
    <w:p> lxml element via pandoc. `src` is markdown; for inline use it must
    already carry its own $...$ delimiters. Returns None on failure."""
    from docx import Document as DocxDoc

    md = f'$$\n{src}\n$$\n' if display else src + '\n'
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        tmp = f.name
    try:
        result = subprocess.run(
            ['pandoc', '-f', 'markdown+tex_math_dollars', '-t', 'docx', '-o', tmp],
            input=md.encode('utf-8'),
            capture_output=True,
        )
        if result.returncode != 0:
            return None
        tmp_doc = DocxDoc(tmp)
        for para in tmp_doc.paragraphs:
            if para._p.find(f'.//{{{_OMML_NS}}}oMath') is not None or para.text.strip():
                return para._p
        return None
    except Exception as exc:
        print(f'  [math] pandoc error: {exc}')
        return None
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass


def _post_process_math(docx_path: str, content_proxies: list) -> None:
    """Replace MATH_DISP_N / MATH_PARA_N sentinels with OMML paragraphs via pandoc."""
    from docx import Document as DocxDoc

    # Build sentinel → (latex_src, is_display, italic, comment_id) map
    math_map: dict[str, tuple[str, bool, bool, int]] = {}
    for proxy in content_proxies:
        if isinstance(proxy, MathDisplayProxy):
            math_map[proxy.math_placeholder] = (proxy._node.math_src, True, False, 0)
        elif isinstance(proxy, HeadingProxy) and proxy.has_math:
            math_map[proxy.math_placeholder] = (
                proxy._node.math_raw, False, False, getattr(proxy, 'comment_id', 0))
        elif isinstance(proxy, FigureProxy) and proxy.has_caption_math:
            math_map[proxy.caption_math_placeholder] = (proxy._caption_text, False, True, 0)

    if not math_map:
        return

    _pandoc_to_para = _pandoc_math_para

    NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _bracket_with_comment(p_el, cid: int) -> None:
        """Wrap the whole OMML paragraph in a commentRangeStart/End pair plus a
        CommentReference run, so a native Word comment can anchor to inline math.
        comments.xml itself is written later by _post_process_comments()."""
        pPr = p_el.find(f'{{{NS_W}}}pPr')
        rng_start = etree.Element(f'{{{NS_W}}}commentRangeStart')
        rng_start.set(f'{{{NS_W}}}id', str(cid))
        if pPr is not None:
            pPr.addnext(rng_start)
        else:
            p_el.insert(0, rng_start)
        rng_end = etree.SubElement(p_el, f'{{{NS_W}}}commentRangeEnd')
        rng_end.set(f'{{{NS_W}}}id', str(cid))
        ref_run = etree.SubElement(p_el, f'{{{NS_W}}}r')
        r_pr = etree.SubElement(ref_run, f'{{{NS_W}}}rPr')
        r_style = etree.SubElement(r_pr, f'{{{NS_W}}}rStyle')
        r_style.set(f'{{{NS_W}}}val', 'CommentReference')
        ref_el = etree.SubElement(ref_run, f'{{{NS_W}}}commentReference')
        ref_el.set(f'{{{NS_W}}}id', str(cid))

    doc = DocxDoc(docx_path)
    replaced = 0
    # Use iter() so paragraphs inside table cells are also found.
    all_para_els = [el for el in doc.element.body.iter()
                    if el.tag == f'{{{NS_W}}}p']
    for p_el in all_para_els:
        txt = ''.join(t.text or '' for t in p_el.findall(f'.//{{{NS_W}}}t')).strip()
        if txt not in math_map:
            continue
        src, display, italic, comment_id = math_map[txt]
        new_p = _pandoc_to_para(src, display)
        if new_p is not None:
            # Apply italic to all text runs when the source is a figure caption.
            if italic:
                for r_el in new_p.findall(f'.//{{{NS_W}}}r'):
                    rPr = r_el.find(f'{{{NS_W}}}rPr')
                    if rPr is None:
                        rPr = etree.Element(f'{{{NS_W}}}rPr')
                        r_el.insert(0, rPr)
                    if rPr.find(f'{{{NS_W}}}i') is None:
                        etree.SubElement(rPr, f'{{{NS_W}}}i')
                    if rPr.find(f'{{{NS_W}}}iCs') is None:
                        etree.SubElement(rPr, f'{{{NS_W}}}iCs')
            # Copy paragraph properties (style, bullet numbering, indent) from
            # the placeholder paragraph so bullet/para formatting is preserved.
            orig_pPr = p_el.find(f'{{{NS_W}}}pPr')
            if orig_pPr is not None:
                new_pPr = new_p.find(f'{{{NS_W}}}pPr')
                orig_copy = copy.deepcopy(orig_pPr)
                if new_pPr is not None:
                    new_p.replace(new_pPr, orig_copy)
                else:
                    new_p.insert(0, orig_copy)
            if comment_id:
                _bracket_with_comment(new_p, comment_id)
            p_el.addnext(new_p)
            p_el.getparent().remove(p_el)
            replaced += 1
        else:
            print(f'  [math] WARNING: failed to convert "{txt[:60]}"')

    if replaced:
        doc.save(docx_path)
        print(f'Post-processed {replaced} math expression(s) into {docx_path}')


# Unescaped $...$ span (non-greedy). A backslash-escaped \$ is not a delimiter,
# so literal-money like \$0.0042 stays text.
_RE_INLINE_MATH_SPAN = re.compile(r'(?<!\\)\$(.+?)(?<!\\)\$', re.S)


def _post_process_inline_math(docx_path: str) -> None:
    """Render inline $...$ spans that survive the sentinel math pass as OMML.

    The sentinel pass (_post_process_math) replaces a whole math *body*
    paragraph with a pandoc-built one, but several render paths emit their text
    verbatim and never go through it: table cells (cell.text = raw markdown),
    table captions ("Table N: ..." strings), and headings. Any $...$ in those
    reaches the DOCX as literal text. This pass walks every paragraph and
    splices each unescaped $...$ span, in place within its text run, into an
    inline <m:oMath>, leaving all surrounding text (and any other markdown it
    carries — **bold**, [brackets], :ref{} sentinels) untouched. Body math
    paragraphs no longer carry a literal $ by this point, so they are skipped
    naturally.
    """
    from docx import Document as DocxDoc
    from docx.oxml.ns import qn

    NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'
    omath_cache: dict[str, object] = {}

    def _omath_for(latex: str):
        """Return a fresh (deep-copied) inline <m:oMath> for the given LaTeX."""
        if latex not in omath_cache:
            para = _pandoc_math_para(f'${latex}$', display=False)
            om = para.find(f'.//{{{_OMML_NS}}}oMath') if para is not None else None
            omath_cache[latex] = om
        cached = omath_cache[latex]
        return copy.deepcopy(cached) if cached is not None else None

    def _text_run(text: str, rPr):
        r = etree.Element(f'{{{NS_W}}}r')
        if rPr is not None:
            r.append(copy.deepcopy(rPr))
        t = etree.SubElement(r, f'{{{NS_W}}}t')
        t.set(XML_SPACE, 'preserve')
        t.text = text
        return r

    doc = DocxDoc(docx_path)
    converted = 0
    for p_el in doc.element.body.iter(qn('w:p')):
        for r_el in list(p_el.findall(qn('w:r'))):
            t_el = r_el.find(qn('w:t'))
            text = t_el.text if t_el is not None else None
            if not text or not _RE_INLINE_MATH_SPAN.search(text):
                continue
            rPr = r_el.find(qn('w:rPr'))
            pos = list(p_el).index(r_el)
            new_nodes = []
            idx = 0
            ok = False
            for m in _RE_INLINE_MATH_SPAN.finditer(text):
                om = _omath_for(m.group(1))
                if om is None:
                    continue  # leave this span as literal text
                if text[idx:m.start()]:
                    new_nodes.append(_text_run(text[idx:m.start()], rPr))
                new_nodes.append(om)
                idx = m.end()
                ok = True
            if not ok:
                continue
            if text[idx:]:
                new_nodes.append(_text_run(text[idx:], rPr))
            for off, node in enumerate(new_nodes):
                p_el.insert(pos + off, node)
            p_el.remove(r_el)
            converted += 1

    if converted:
        doc.save(docx_path)
        print(f'Post-processed inline math in {converted} run(s) into {docx_path}')


def _inline_to_anchor(inline_el) -> None:
    """Convert a <wp:inline> element in-place to <wp:anchor> with Top-and-Bottom wrapping.

    Modifies the element's tag and attributes, inserts the required positional
    child elements, and adds <wp:wrapTopAndBottom/>.  The graphic payload
    (extent, effectExtent, docPr, cNvGraphicFramePr, a:graphic) is preserved.
    """
    WP = _NS_WP

    # Change tag from wp:inline → wp:anchor
    inline_el.tag = f'{{{WP}}}anchor'

    # Required anchor attributes (inline has distT/distB/distL/distR already
    # but lacks the rest)
    inline_el.set('simplePos', '0')
    inline_el.set('relativeHeight', '251658240')
    inline_el.set('behindDoc', '0')
    inline_el.set('locked', '0')
    inline_el.set('layoutInCell', '1')
    inline_el.set('allowOverlap', '0')
    for dist in ('distT', 'distB', 'distL', 'distR'):
        if inline_el.get(dist) is None:
            inline_el.set(dist, '0')

    # Build positional elements to insert at the front (before wp:extent)
    simple_pos = etree.Element(f'{{{WP}}}simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')

    pos_h = etree.Element(f'{{{WP}}}positionH')
    pos_h.set('relativeFrom', 'column')
    align_h = etree.SubElement(pos_h, f'{{{WP}}}align')
    align_h.text = 'center'

    pos_v = etree.Element(f'{{{WP}}}positionV')
    pos_v.set('relativeFrom', 'paragraph')
    align_v = etree.SubElement(pos_v, f'{{{WP}}}align')
    align_v.text = 'top'

    # Insert positional elements before the first child (wp:extent)
    extent_el = inline_el.find(f'{{{WP}}}extent')
    idx = list(inline_el).index(extent_el) if extent_el is not None else 0
    for i, el in enumerate([simple_pos, pos_h, pos_v]):
        inline_el.insert(idx + i, el)

    # Insert <wp:wrapTopAndBottom/> after wp:extent (and wp:effectExtent if present)
    effect_el = inline_el.find(f'{{{WP}}}effectExtent')
    anchor_child = effect_el if effect_el is not None else inline_el.find(f'{{{WP}}}extent')
    if anchor_child is not None:
        wrap_el = etree.Element(f'{{{WP}}}wrapTopAndBottom')
        anchor_child.addnext(wrap_el)
    else:
        inline_el.append(etree.Element(f'{{{WP}}}wrapTopAndBottom'))


_RE_SECLABEL = re.compile(r'@@SECLABEL:([^@]+)@@')
_RE_SECREF_SENT = re.compile(r'@@SECREF:([^@]+)@@')
_RE_FIGREF_SENT = re.compile(r'@@FIGREF:([^@]+)@@')
_RE_TABREF_SENT = re.compile(r'@@TABREF:([^@]+)@@')


def _post_process_figures(docx_path: str, figure_title_style: str = 'FImageTitle') -> None:
    """Wrap each figure block in a no-border 2-row table (image row + caption row).

    Finds FIGURE_START_N / FIGURE_END_N sentinel paragraphs placed by the Jinja
    template, collects the image and caption paragraphs between them, and replaces
    the whole group with a borderless table:

      Row 0  — image (centered, keepNext so it stays with Row 1)
      Row 1  — figure-title paragraph (bold italic) + caption paragraph(s)

    figure_title_style is the template's style ID for figure title paragraphs
    (a RENDER_OPTIONS key exported by the project's template prep script).

    Spacing of 160 twips is added above Row 0 and below the last caption paragraph
    to separate the figure from surrounding body text.  Both rows have cantSplit so
    the table does not break across pages.

    Must run BEFORE _post_process_math so that MATH_CAP_N sentinels, if present
    in the caption paragraph, are still found by doc.paragraphs after being moved
    into the table cell.
    """
    from docx import Document as DocxDoc
    from docx.oxml import OxmlElement

    NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

    def _get_text(el):
        return ''.join(t.text or '' for t in el.findall(f'.//{{{NS_W}}}t'))

    def _set_spacing(p_el, before=None, after=None, line=None, line_rule=None):
        pPr = p_el.find(f'{{{NS_W}}}pPr')
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_el.insert(0, pPr)
        spacing = pPr.find(f'{{{NS_W}}}spacing')
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        if before is not None:
            spacing.set(f'{{{NS_W}}}before', str(before))
        if after is not None:
            spacing.set(f'{{{NS_W}}}after', str(after))
        if line is not None:
            spacing.set(f'{{{NS_W}}}line', str(line))
        if line_rule is not None:
            spacing.set(f'{{{NS_W}}}lineRule', line_rule)

    def _set_keep_next(p_el):
        pPr = p_el.find(f'{{{NS_W}}}pPr')
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_el.insert(0, pPr)
        if pPr.find(f'{{{NS_W}}}keepNext') is None:
            pPr.append(OxmlElement('w:keepNext'))

    def _make_cant_split_tr_pr():
        trPr = OxmlElement('w:trPr')
        cs = OxmlElement('w:cantSplit')
        trPr.append(cs)
        return trPr

    def _make_auto_tc(paras):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(f'{{{NS_W}}}w', '0')
        tcW.set(f'{{{NS_W}}}type', 'auto')
        tcPr.append(tcW)
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'left', 'bottom', 'right'):
            m = OxmlElement(f'w:{side}')
            m.set(f'{{{NS_W}}}w', '120')
            m.set(f'{{{NS_W}}}type', 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        tc.append(tcPr)
        for p in paras:
            tc.append(p)
        return tc

    def _build_table(image_paras, caption_paras):
        tbl = OxmlElement('w:tbl')

        # Table properties
        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(f'{{{NS_W}}}w', '5000')
        tblW.set(f'{{{NS_W}}}type', 'pct')
        tblPr.append(tblW)
        tblBorders = OxmlElement('w:tblBorders')
        for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{name}')
            if name in ('top', 'bottom'):
                b.set(f'{{{NS_W}}}val', 'single')
                b.set(f'{{{NS_W}}}sz', '6')
                b.set(f'{{{NS_W}}}space', '0')
                b.set(f'{{{NS_W}}}color', 'auto')
            else:
                b.set(f'{{{NS_W}}}val', 'none')
                b.set(f'{{{NS_W}}}sz', '0')
                b.set(f'{{{NS_W}}}space', '0')
                b.set(f'{{{NS_W}}}color', 'auto')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        # Single-column grid
        tblGrid = OxmlElement('w:tblGrid')
        tblGrid.append(OxmlElement('w:gridCol'))
        tbl.append(tblGrid)

        # Row 0: image
        tr0 = OxmlElement('w:tr')
        tr0.append(_make_cant_split_tr_pr())
        tr0.append(_make_auto_tc(image_paras))
        tbl.append(tr0)

        # Row 1: caption (title + long description)
        tr1 = OxmlElement('w:tr')
        tr1.append(_make_cant_split_tr_pr())
        tr1.append(_make_auto_tc(caption_paras))
        tbl.append(tr1)

        return tbl

    doc = DocxDoc(docx_path)
    body = doc.element.body
    found = 0

    while True:
        children = list(body)
        start_el = None
        fig_num = None
        for el in children:
            if el.tag != f'{{{NS_W}}}p':
                continue
            txt = _get_text(el).strip()
            if txt.startswith('FIGURE_START_'):
                try:
                    fig_num = int(txt[len('FIGURE_START_'):])
                    start_el = el
                    break
                except ValueError:
                    continue
        if start_el is None:
            break

        end_sentinel = f'FIGURE_END_{fig_num}'
        content_els = []
        end_el = None
        collecting = False
        for el in children:
            if el is start_el:
                collecting = True
                continue
            if not collecting:
                continue
            if _get_text(el).strip() == end_sentinel:
                end_el = el
                break
            content_els.append(el)

        if end_el is None:
            body.remove(start_el)
            continue

        # Split into image paragraphs vs caption paragraphs
        image_paras, caption_paras = [], []
        for el in content_els:
            if el.tag == f'{{{NS_W}}}p' and (
                el.find(f'.//{{{NS_WP}}}inline') is not None or
                el.find(f'.//{{{NS_WP}}}anchor') is not None
            ):
                image_paras.append(el)
            else:
                caption_paras.append(el)

        # Bold the runs in any figure-title paragraph in the caption cell
        for p in caption_paras:
            if p.tag != f'{{{NS_W}}}p':
                continue
            ps = p.find(f'.//{{{NS_W}}}pStyle')
            if ps is None or ps.get(f'{{{NS_W}}}val') != figure_title_style:
                continue
            for r in p.findall(f'{{{NS_W}}}r'):
                rPr = r.find(f'{{{NS_W}}}rPr')
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                for tag in ('b', 'bCs'):
                    if rPr.find(f'{{{NS_W}}}' + tag) is None:
                        rPr.append(OxmlElement(f'w:{tag}'))

        # Spacing: 160 twips above first image para, zero after each image para so
        # Normal-style paragraph spacing does not bleed into the caption row below.
        # line=1/lineRule=exact collapses the anchor paragraph's own line height to
        # 1 twip so the phantom empty line below the image disappears.
        if image_paras:
            _set_spacing(image_paras[0], before=160, after=0, line=1, line_rule='exact')
            for p in image_paras[1:]:
                _set_spacing(p, after=0, line=1, line_rule='exact')
            for p in image_paras:
                _set_keep_next(p)
        if caption_paras:
            _set_spacing(caption_paras[-1], after=160)

        # Fallback empty paragraphs
        if not image_paras:
            image_paras = [OxmlElement('w:p')]
        if not caption_paras:
            caption_paras = [OxmlElement('w:p')]

        # Detach content elements from body
        for el in content_els:
            body.remove(el)

        # Build table with the detached elements and insert before start sentinel
        tbl = _build_table(image_paras, caption_paras)
        start_el.addprevious(tbl)

        # Remove the two sentinel paragraphs
        body.remove(start_el)
        body.remove(end_el)

        # Insert a zero-height empty paragraph after the table so adjacent
        # figure tables are not merged by Word into a single table.
        spacer = OxmlElement('w:p')
        spacer_pPr = OxmlElement('w:pPr')
        spacer_spacing = OxmlElement('w:spacing')
        spacer_spacing.set(f'{{{NS_W}}}before', '0')
        spacer_spacing.set(f'{{{NS_W}}}after', '0')
        spacer_spacing.set(f'{{{NS_W}}}line', '1')
        spacer_spacing.set(f'{{{NS_W}}}lineRule', 'exact')
        spacer_pPr.append(spacer_spacing)
        spacer.append(spacer_pPr)
        tbl.addnext(spacer)

        found += 1

    if found:
        doc.save(docx_path)
        print(f'Post-processed {found} figure(s) into borderless 2-row tables in {docx_path}')


def _post_process_crossrefs(
    docx_path: str,
    content_nodes: list[ContentNode],
    fig_map: dict[str, int],
    tab_map: dict[str, int],
    figure_title_style: str = 'FImageTitle',
) -> None:
    """Insert Word bookmarks on labelled headings, figure titles, and REF/HYPERLINK fields at cross-reference sites."""
    from docx import Document as DocxDoc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    doc = DocxDoc(docx_path)

    # Start bookmark IDs above any already present in the document
    existing_ids = [
        int(bm.get(qn('w:id'), '0'))
        for bm in doc.element.body.findall(f'.//{{{W}}}bookmarkStart')
    ]
    bookmark_counter = [max(existing_ids) + 1 if existing_ids else 0]

    def _safe_bm_name(sec_id: str) -> str:
        """Word bookmark names: letters, digits, underscores only, start with letter."""
        return 'sec_' + re.sub(r'[^A-Za-z0-9_]', '_', sec_id)

    def _safe_fig_bm_name(fig_id: str) -> str:
        """Bookmark name for a figure."""
        return 'fig_' + re.sub(r'[^A-Za-z0-9_]', '_', fig_id)

    def _safe_tab_bm_name(tab_id: str) -> str:
        """Bookmark name for a table."""
        return 'tab_' + re.sub(r'[^A-Za-z0-9_]', '_', tab_id)

    def _fld_run(fld_type: str) -> 'etree._Element':
        r = OxmlElement('w:r')
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), fld_type)
        r.append(fc)
        return r

    def _instr_run(instr: str) -> 'etree._Element':
        r = OxmlElement('w:r')
        t = OxmlElement('w:instrText')
        t.set(XML_SPACE, 'preserve')
        t.text = instr
        r.append(t)
        return r

    def _text_run(text: str, source_run: 'etree._Element | None' = None) -> 'etree._Element':
        r = OxmlElement('w:r')
        if source_run is not None:
            src_rpr = source_run.find(f'{{{W}}}rPr')
            if src_rpr is not None:
                r.append(copy.deepcopy(src_rpr))
        t = OxmlElement('w:t')
        t.set(XML_SPACE, 'preserve')
        t.text = text
        r.append(t)
        return r

    replaced_labels = 0
    replaced_refs = 0
    replaced_fig_refs = 0
    replaced_tab_refs = 0

    all_paras = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    # --- Phase 1: Create bookmarks on figure title paragraphs (figure_title_style) ---
    # Use raw XML iteration because figure tables are built with lxml elements
    # that python-docx's doc.tables / doc.paragraphs don't enumerate.
    fig_num_to_id = {v: k for k, v in fig_map.items()}

    all_p_els = list(doc.element.body.iter(f'{{{W}}}p'))

    # --- Phase 0: Repair malformed nested runs (<w:r><w:t><w:r>...</w:r></w:t></w:r>) ---
    # docxtpl renders a RichText caption via plain {{ }} (not {{r }}), nesting the
    # formatted runs inside the placeholder run's <w:t>.  This is invalid OOXML and
    # hides any sentinel inside the nested run from the cross-reference scanner.
    # Lift the nested runs out so they become valid sibling runs of the paragraph.
    for p in all_p_els:
        for outer in list(p.findall(f'{{{W}}}r')):
            t_el = outer.find(f'{{{W}}}t')
            if t_el is None:
                continue
            nested = t_el.findall(f'{{{W}}}r')
            if not nested:
                continue
            outer_rpr = outer.find(f'{{{W}}}rPr')
            new_runs: list['etree._Element'] = []

            def _wrap_text(text: str) -> 'etree._Element':
                r = OxmlElement('w:r')
                if outer_rpr is not None:
                    r.append(copy.deepcopy(outer_rpr))
                t = OxmlElement('w:t')
                t.set(XML_SPACE, 'preserve')
                t.text = text
                r.append(t)
                return r

            if t_el.text:
                new_runs.append(_wrap_text(t_el.text))
            for inner in nested:
                tail = inner.tail
                inner.tail = None
                new_runs.append(inner)
                if tail:
                    new_runs.append(_wrap_text(tail))

            for nr in new_runs:
                outer.addprevious(nr)
            p.remove(outer)

    for p in all_p_els:
        pPr = p.find(f'{{{W}}}pPr')
        ps = pPr.find(f'{{{W}}}pStyle') if pPr is not None else None
        if ps is None or ps.get(f'{{{W}}}val') != figure_title_style:
            continue
        txt = ''.join(t.text or '' for t in p.findall(f'.//{{{W}}}t')).strip()
        m = re.match(r'Figure\s+(\d+)', txt)
        if not m:
            continue
        fig_num = int(m.group(1))
        fig_id = fig_num_to_id.get(fig_num)
        if not fig_id:
            continue
        bm_name = _safe_fig_bm_name(fig_id)
        bid = bookmark_counter[0]
        bookmark_counter[0] += 1

        first_run = p.find(f'{{{W}}}r')
        if first_run is not None:
            bm_start = OxmlElement('w:bookmarkStart')
            bm_start.set(qn('w:id'), str(bid))
            bm_start.set(qn('w:name'), bm_name)
            first_run.addprevious(bm_start)

            bm_end = OxmlElement('w:bookmarkEnd')
            bm_end.set(qn('w:id'), str(bid))
            first_run.addnext(bm_end)

    # --- Phase 1b: Create bookmarks on table caption paragraphs ---
    tab_num_to_id = {v: k for k, v in tab_map.items()}

    for para in all_paras:
        p = para._p
        txt = para.text.strip()
        m = re.match(r'Table\s+(\d+)', txt)
        if not m:
            continue
        tab_num = int(m.group(1))
        tab_id = tab_num_to_id.get(tab_num)
        if not tab_id:
            continue
        bm_name = _safe_tab_bm_name(tab_id)
        bid = bookmark_counter[0]
        bookmark_counter[0] += 1

        first_run = p.find(f'{{{W}}}r')
        if first_run is not None:
            bm_start = OxmlElement('w:bookmarkStart')
            bm_start.set(qn('w:id'), str(bid))
            bm_start.set(qn('w:name'), bm_name)
            first_run.addprevious(bm_start)

            bm_end = OxmlElement('w:bookmarkEnd')
            bm_end.set(qn('w:id'), str(bid))
            first_run.addnext(bm_end)

    # --- Phase 1c: Repair sentinels split across runs by pandoc math rendering ---
    # Pandoc may split @@SECREF:id@@ / @@FIGREF:id@@ / @@TABREF:id@@ /
    # @@FOOTREF:key@@ across adjacent w:r elements (e.g. @ | @FIGREF:id | @@).
    # A paragraph can contain several such sentinels back-to-back, so keep
    # repairing until none remain. FOOTREF sentinels are repaired here and
    # consumed later by _post_process_footnotes.
    _RE_ANY_SENT = re.compile(r'@@(?:SECREF|FIGREF|TABREF|FOOTREF):[^@]+@@')
    for para in all_paras:
        p = para._p
        while True:
            runs = list(p.findall(f'{{{W}}}r'))
            segments: list[tuple[int, int, 'etree._Element', str]] = []
            pos = 0
            for r in runs:
                rtext = ''.join(t.text or '' for t in r.findall(f'{{{W}}}t'))
                segments.append((pos, pos + len(rtext), r, rtext))
                pos += len(rtext)
            combined = ''.join(s[3] for s in segments)

            broken = None
            for m in _RE_ANY_SENT.finditer(combined):
                spanning = [s for s in segments if s[0] < m.end() and s[1] > m.start()]
                if len(spanning) > 1:
                    broken = (m, spanning)
                    break
            if broken is None:
                break

            m, spanning = broken
            target_runs = [s[2] for s in spanning]
            merged_text = ''.join(s[3] for s in spanning)
            m2 = _RE_ANY_SENT.search(merged_text)
            if not m2:
                break
            prefix = merged_text[:m2.start()]
            suffix = merged_text[m2.end():]
            # Build replacement runs: prefix (optional), sentinel run, suffix (optional)
            new_runs: list['etree._Element'] = []
            if prefix:
                pr = OxmlElement('w:r')
                pt = OxmlElement('w:t')
                pt.set(XML_SPACE, 'preserve')
                pt.text = prefix
                pr.append(pt)
                new_runs.append(pr)
            sr = copy.deepcopy(target_runs[0])
            for st in sr.findall(f'{{{W}}}t'):
                st.text = m2.group(0)
            new_runs.append(sr)
            if suffix:
                su = OxmlElement('w:r')
                sut = OxmlElement('w:t')
                sut.set(XML_SPACE, 'preserve')
                sut.text = suffix
                su.append(sut)
                new_runs.append(su)
            # Anchor insertion at the run immediately following the split block
            # (positionally correct for mid-paragraph splits, not just paragraph-initial ones).
            last_target_idx = max(runs.index(r) for r in target_runs)
            next_run = runs[last_target_idx + 1] if last_target_idx + 1 < len(runs) else None
            for r in target_runs:
                p.remove(r)
            if next_run is not None:
                # addprevious on a fixed anchor: each call lands immediately
                # before it, so iterate in forward order to keep new_runs in sequence.
                for nr in new_runs:
                    next_run.addprevious(nr)
            else:
                for nr in new_runs:
                    p.append(nr)

    # --- Phase 2: Process cross-reference sentinels in all paragraphs ---
    # Iterate raw w:p elements so paragraphs inside lxml-built figure tables
    # (not enumerated by python-docx's doc.tables) are also processed.
    for p in all_p_els:
        run_queue = collections.deque(p.findall(f'{{{W}}}r'))
        while run_queue:
            run = run_queue.popleft()
            for t_el in run.findall(f'{{{W}}}t'):
                txt = t_el.text or ''

                # --- Heading bookmark: @@SECLABEL:id@@ ---
                lm = _RE_SECLABEL.search(txt)
                if lm:
                    sec_id = lm.group(1)
                    bm_name = _safe_bm_name(sec_id)
                    bid = bookmark_counter[0]
                    bookmark_counter[0] += 1

                    # Strip sentinel from the run text
                    t_el.text = txt[:lm.start()] + txt[lm.end():]

                    bm_start = OxmlElement('w:bookmarkStart')
                    bm_start.set(qn('w:id'), str(bid))
                    bm_start.set(qn('w:name'), bm_name)
                    run.addprevious(bm_start)

                    bm_end = OxmlElement('w:bookmarkEnd')
                    bm_end.set(qn('w:id'), str(bid))
                    run.addnext(bm_end)

                    replaced_labels += 1
                    continue  # a run carries at most one sentinel

                # --- Section cross-reference field: @@SECREF:id@@ ---
                rm = _RE_SECREF_SENT.search(txt)
                if rm:
                    ref_id = rm.group(1)
                    bm_name = _safe_bm_name(ref_id)

                    before_txt = txt[:rm.start()]
                    after_txt  = txt[rm.end():]

                    # Emit a native Word REF field with \n (paragraph number)
                    # and \h (hyperlink).  "Section " is static text before the
                    # field; the REF generates the heading number from Word's
                    # heading style numbering.  No custom numbering needed.
                    tail_elements = [
                        _text_run('Section ', source_run=run),
                        _fld_run('begin'),
                        _instr_run(f' REF "{bm_name}" \\n \\h '),
                        _fld_run('separate'),
                        _text_run('?', source_run=run),
                        _fld_run('end'),
                    ]
                    after_run = None
                    if after_txt:
                        after_run = _text_run(after_txt, source_run=run)
                        tail_elements.append(after_run)

                    for el in reversed(tail_elements):
                        run.addnext(el)

                    if after_run is not None:
                        run_queue.appendleft(after_run)  # process remaining sentinels

                    # Keep before-text in the current run, or clear it if empty
                    t_el.text = before_txt if before_txt else ''

                    replaced_refs += 1
                    continue

                # --- Figure cross-reference field: @@FIGREF:id@@ ---
                fm = _RE_FIGREF_SENT.search(txt)
                if fm:
                    fig_id = fm.group(1)
                    bm_name = _safe_fig_bm_name(fig_id)
                    fig_num = fig_map.get(fig_id)
                    display = f'Figure {fig_num}' if fig_num else f'[?{fig_id}]'

                    before_txt = txt[:fm.start()]
                    after_txt  = txt[fm.end():]

                    # REF field with \h switch (hyperlink) and \* MERGEFORMAT.
                    # The REF \h bookmark creates a clickable hyperlink to the figure.
                    # We use a static display string so F9 doesn't change it.
                    tail_elements = [
                        _fld_run('begin'),
                        _instr_run(f' REF "{bm_name}" \\h'),
                        _fld_run('separate'),
                        _text_run(display, source_run=run),
                        _fld_run('end'),
                    ]
                    after_run = None
                    if after_txt:
                        after_run = _text_run(after_txt, source_run=run)
                        tail_elements.append(after_run)

                    for el in reversed(tail_elements):
                        run.addnext(el)

                    if after_run is not None:
                        run_queue.appendleft(after_run)

                    if before_txt:
                        t_el.text = before_txt
                    else:
                        p.remove(run)

                    replaced_fig_refs += 1

                # --- Table cross-reference field: @@TABREF:id@@ ---
                tm = _RE_TABREF_SENT.search(txt)
                if tm:
                    tab_id = tm.group(1)
                    bm_name = _safe_tab_bm_name(tab_id)
                    tab_num = tab_map.get(tab_id)
                    display = f'Table {tab_num}' if tab_num else f'[?{tab_id}]'

                    before_txt = txt[:tm.start()]
                    after_txt  = txt[tm.end():]

                    tail_elements = [
                        _fld_run('begin'),
                        _instr_run(f' REF "{bm_name}" \\h'),
                        _fld_run('separate'),
                        _text_run(display, source_run=run),
                        _fld_run('end'),
                    ]
                    after_run = None
                    if after_txt:
                        after_run = _text_run(after_txt, source_run=run)
                        tail_elements.append(after_run)

                    for el in reversed(tail_elements):
                        run.addnext(el)

                    if after_run is not None:
                        run_queue.appendleft(after_run)

                    if before_txt:
                        t_el.text = before_txt
                    else:
                        p.remove(run)

                    replaced_tab_refs += 1

    if replaced_labels or replaced_refs or replaced_fig_refs or replaced_tab_refs:
        doc.save(docx_path)
        print(f'  Post-processed {replaced_labels} bookmark(s), {replaced_refs} section ref(s), '
              f'{replaced_fig_refs} figure ref(s), {replaced_tab_refs} table ref(s)')


def _post_process_image_wrapping(docx_path: str) -> None:
    """Replace wp:inline image elements with wp:anchor (Top and Bottom wrapping)."""
    from docx import Document as DocxDoc

    doc = DocxDoc(docx_path)
    WP = _NS_WP
    inlines = list(doc.element.body.findall(f'.//{{{WP}}}inline'))
    if not inlines:
        return
    for el in inlines:
        _inline_to_anchor(el)
    doc.save(docx_path)
    print(f'Post-processed {len(inlines)} image(s) → Top and Bottom wrapping in {docx_path}')


def _validate_template_context(tpl: DocxTemplate, context: dict) -> None:
    """Validate that template Jinja variables have corresponding context keys.

    Uses docxtpl's get_undeclared_template_variables() (jinja2's own AST
    analysis), which sees every variable reference — including ones split
    across Word runs or used in filters/conditions — and correctly excludes
    loop-local variables.  Checks both directions against the context dict.

    Raises SystemExit on mismatch.
    """
    template_vars = tpl.get_undeclared_template_variables()
    context_keys = set(context.keys())

    errors = []

    # 1. Every template variable must have a context key
    for var in sorted(template_vars):
        if var not in context_keys:
            errors.append(f'Template variable "{var}" not found in context')

    # 2. Every context key must be used by the template
    for key in sorted(context_keys):
        if key not in template_vars:
            errors.append(f'Context key "{key}" not used in template')

    if errors:
        print('\nTemplate-context mismatch:')
        for e in errors:
            print(f'  {e}')
        raise SystemExit(1)


def render(
    template_path: str,
    context: dict,
    body_var_name: str | list[str],
    output_path: str,
    image_base: str = '.',
    figure_title_style: str = 'FImageTitle',
) -> None:
    """Render the instrumented template with a context dict and write to output_path.

    Args:
        template_path: Path to the Jinja-instrumented DOCX (from the project's
                       template prep script).
        context:       Context dict (from the VAR blocks + parsed body).
        body_var_name: Key(s) in context holding ContentNode lists. A single
                       name for the common one-freeform-VAR template; a list
                       of names for a fixed-structure template with multiple
                       independent content loops (see the jinjify-template
                       skill). All lists are concatenated before figure/table
                       numbering and math/inline-run sentinel assignment so
                       numbering and cross-references stay consistent across
                       the whole document, then split back apart so each VAR
                       renders through its own template loop.
        output_path:   Destination DOCX file.
        image_base:    Directory to search for figure image files.
        figure_title_style:
                       Template style ID for figure title paragraphs; the
                       figure/cross-ref post-processing keys on it.  Projects
                       override it via RENDER_OPTIONS in their prep script.
    """
    tpl = DocxTemplate(template_path)
    img_base = Path(image_base)

    sec = Document(template_path).sections[0]
    text_width = sec.page_width - sec.left_margin - sec.right_margin

    body_var_names = [body_var_name] if isinstance(body_var_name, str) else list(body_var_name)

    # Concatenate every freeform VAR's content nodes so figure/table numbers
    # and math/inline sentinels are assigned once, globally, in draft order —
    # not reset to zero (and colliding) at each VAR's boundary.
    content_nodes: list[ContentNode] = []
    boundaries: list[tuple[str, int]] = []
    for name in body_var_names:
        nodes = context[name]
        content_nodes.extend(nodes)
        boundaries.append((name, len(nodes)))

    # Finalize: assign sequential figure/table numbers, build maps
    figures_list, tables_list, fig_map, tab_map = finalize_content(content_nodes)

    # Footnotes: registry in first-reference order (validates defs/refs)
    footnotes = _collect_footnotes(content_nodes)

    # Build content proxy list (assigns math counters)
    content_proxies, inline_map, comments = _build_content_proxies(content_nodes, tpl, img_base, text_width)

    # Split proxies back into each VAR's own list, in the template's context
    offset = 0
    for name, length in boundaries:
        context[name] = content_proxies[offset:offset + length]
        offset += length

    # Structured (non-body) VARs: markdown-image strings → native InlineImage
    body_set = set(body_var_names)
    for key in context:
        if key not in body_set:
            context[key] = _convert_context_images(context[key], tpl, img_base, text_width)

    _validate_template_context(tpl, context)

    tpl.render(context)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    tpl.save(output_path)

    # Post-process: replace TABLE_PLACEHOLDER_N paragraphs with real tables
    _post_process_tables(output_path, tables_list)
    # Post-process: restore inline bold/italic runs for headings/paragraphs/bullets
    _post_process_inline_runs(output_path, inline_map)
    # Post-process: wrap figure blocks in borderless 2-row tables
    _post_process_figures(output_path, figure_title_style)
    # Post-process: replace MATH_DISP_N / MATH_PARA_N sentinels with OMML
    _post_process_math(output_path, content_proxies)
    # Post-process: render inline $...$ spans left verbatim by non-body paths
    # (table cells, table captions, headings)
    _post_process_inline_math(output_path)
    # Post-process: insert Word bookmarks and REF fields for section cross-references
    _post_process_crossrefs(output_path, content_nodes, fig_map, tab_map,
                            figure_title_style)
    # Post-process: convert inline images to Top and Bottom wrapping
    _post_process_image_wrapping(output_path)
    # Post-process: sync native w14:checkbox controls from @@CHK:x@@ sentinels
    _post_process_checkboxes(output_path)
    # Post-process: convert @@FOOTREF:key@@ sentinels into native Word
    # footnotes (footnotes.xml entries + footnoteReference runs).
    _post_process_footnotes(output_path, footnotes)
    # Post-process: convert @@COMMENTSTART/END:n@@ sentinels into native Word
    # comments. Must run last (see _post_process_comments docstring).
    _post_process_comments(output_path, comments)
    print(f'Rendered → {output_path}')
